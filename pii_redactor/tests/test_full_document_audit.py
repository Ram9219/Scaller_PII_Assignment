import os
import sys
import docx
import pytest

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from document.processor import DocumentProcessor
from detectors import RegexDetector, NERDetector, ContextDetector
from core.resolver import EntityResolver
from core.replacement import DeterministicReplacer
from config import FAKER_SEED
from evaluation.audit_redacted_document import audit

def _create_mock_docx(filename: str):
    doc = docx.Document()
    
    # Main paragraph
    doc.add_paragraph("KMPs including, Amit Kumar, Rajesh Sharma and Priya Mehta are also our Executive Directors.")
    
    # Table with nested table (python-docx doesn't easily create nested tables via API, but we can test normal tables)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Email"
    table.cell(0, 1).text = "Phone"
    table.cell(1, 0).text = "test@example.com"
    table.cell(1, 1).text = "9876543210"
    
    # Header Table
    section = doc.sections[0]
    header = section.header
    htable = header.add_table(rows=1, cols=1, width=5000)
    htable.cell(0, 0).text = "CONFIDENTIAL: Contact Sarthak Malvadkar at sarthak@example.com"
    
    doc.save(filename)
    return filename

@pytest.fixture
def processor():
    detectors = [RegexDetector(), NERDetector(), ContextDetector()]
    resolver = EntityResolver()
    replacer = DeterministicReplacer(seed=FAKER_SEED)
    return DocumentProcessor(detectors, resolver, replacer)

def test_full_document_audit(processor, tmp_path):
    orig_file = str(tmp_path / "orig.docx")
    redacted_file = str(tmp_path / "redacted.docx")
    
    _create_mock_docx(orig_file)
    
    # Process
    processor.process(orig_file, redacted_file)
    
    # Audit
    report = audit(orig_file, redacted_file)
    
    assert report["total_original_pii_candidates"] > 0
    assert len(report["surviving_pii"]) == 0
    assert report["coverage"] == 1.0
