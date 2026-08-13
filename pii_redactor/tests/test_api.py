import os
import pytest
from fastapi.testclient import TestClient
from api import app
import docx
import tempfile

client = TestClient(app)

def test_health_check():
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok", "service": "PII Redactor"}

def test_redact_missing_file():
    response = client.post("/redact")
    assert response.status_code == 422 # FastAPI default for missing required field

def test_redact_invalid_file_type():
    # Create a temporary txt file
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
        f.write(b"Hello world")
        tmp_path = f.name
        
    with open(tmp_path, "rb") as f:
        response = client.post("/redact", files={"file": ("test.txt", f, "text/plain")})
        
    os.remove(tmp_path)
    assert response.status_code == 415
    assert "Unsupported file type" in response.json()["detail"]

def test_redact_valid_docx():
    # Create a simple valid docx file
    doc = docx.Document()
    doc.add_paragraph("Hello John Doe, your phone number is 9876543210.")
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        tmp_path = f.name
    
    doc.save(tmp_path)
    
    # Track original structure
    orig_paragraphs = len(doc.paragraphs)
    
    # Upload and redact
    with open(tmp_path, "rb") as f:
        response = client.post("/redact", files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
        
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    # Save the response to a temporary file to verify its contents
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out_f:
        out_f.write(response.content)
        out_path = out_f.name
        
    try:
        # Verify it's a valid DOCX and structural integrity is maintained
        redacted_doc = docx.Document(out_path)
        assert len(redacted_doc.paragraphs) == orig_paragraphs
        
        # Verify redaction occurred
        text = redacted_doc.paragraphs[0].text
        assert "John Doe" not in text
        assert "9876543210" not in text
    finally:
        os.remove(tmp_path)
        os.remove(out_path)
