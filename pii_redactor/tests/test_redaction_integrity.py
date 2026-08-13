import os
import sys
import pytest
import docx

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from detectors import RegexDetector, NERDetector, ContextDetector
from core.resolver import EntityResolver
from core.replacement import DeterministicReplacer
from document.processor import DocumentProcessor
from config import FAKER_SEED

def _create_mock_docx(filename: str, texts: list):
    doc = docx.Document()
    for text in texts:
        doc.add_paragraph(text)
    doc.save(filename)
    return filename

@pytest.fixture
def processor():
    detectors = [RegexDetector(), NERDetector(), ContextDetector()]
    resolver = EntityResolver()
    replacer = DeterministicReplacer(seed=FAKER_SEED)
    return DocumentProcessor(detectors, resolver, replacer)

def test_non_pii_preservation(processor, tmp_path):
    """
    Ensure domain terminology is NOT redacted.
    """
    input_file = str(tmp_path / "input.docx")
    output_file = str(tmp_path / "output.docx")
    
    domain_texts = [
        "The Offer comprises of a Fresh Issue and an Offer for Sale.",
        "Equity Shares are allocated to Anchor Investors.",
        "The Book Building Process is regulated by SEBI ICDR Regulations.",
        "According to the Companies Act, the Stock Exchanges will list the shares.",
        "Categories include QIB, NII, and RII.",
        "The financial amounts reached ₹ 1,200 million or 10%.",
        "See page 45 for more details."
    ]
    
    _create_mock_docx(input_file, domain_texts)
    
    processor.process(input_file, output_file)
    
    doc_out = docx.Document(output_file)
    
    for i, expected_text in enumerate(domain_texts):
        assert doc_out.paragraphs[i].text == expected_text, f"Failed preservation on: {expected_text}"

def test_span_safety(processor, tmp_path):
    """
    Ensure only the PII span is changed and surrounding text remains unchanged.
    """
    input_file = str(tmp_path / "input.docx")
    output_file = str(tmp_path / "output.docx")
    
    # "Sarthak Malvadkar" should be replaced, but "Contact Person:" and ", Company Secretary" must remain exactly.
    test_text = "Contact Person: Sarthak Malvadkar, Company Secretary"
    
    _create_mock_docx(input_file, [test_text])
    
    processor.process(input_file, output_file)
    
    doc_out = docx.Document(output_file)
    out_text = doc_out.paragraphs[0].text
    
    assert "Contact Person: " in out_text
    assert ", Company Secretary" in out_text
    assert "Sarthak Malvadkar" not in out_text

def test_structural_preservation(processor, tmp_path):
    """
    Basic structural check: para counts, table counts.
    (Detailed structural check is handled in test_docx_preservation, but we re-verify here).
    """
    input_file = str(tmp_path / "input.docx")
    output_file = str(tmp_path / "output.docx")
    
    doc = docx.Document()
    doc.add_paragraph("Paragraph 1")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Email"
    table.cell(0, 1).text = "Phone"
    table.cell(1, 0).text = "test@example.com"
    table.cell(1, 1).text = "9876543210"
    doc.save(input_file)
    
    processor.process(input_file, output_file)
    
    doc_out = docx.Document(output_file)
    assert len(doc_out.paragraphs) == 1
    assert len(doc_out.tables) == 1
    assert len(doc_out.tables[0].rows) == 2
    assert len(doc_out.tables[0].columns) == 2
    
    # Ensure PII was removed
    assert "test@example.com" not in doc_out.tables[0].cell(1, 0).text
