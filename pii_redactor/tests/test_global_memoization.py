import pytest
import docx
import os
from document.processor import DocumentProcessor
from detectors import RegexDetector, NERDetector, ContextDetector
from core.resolver import EntityResolver
from core.replacement import DeterministicReplacer
from evaluation.audit_redacted_document import audit, get_expected_text

def test_global_memoization_paragraph_to_table(tmpdir):
    doc = docx.Document()
    # Paragraph 1 provides strong PERSON context
    doc.add_paragraph("Alice Bobson and Charlie Doe are the promoters of our Company.")
    # Table has no context, normally NER might fail or ContextDetector would reject
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Alice Bobson"
    
    input_path = str(tmpdir.join("test_in.docx"))
    output_path = str(tmpdir.join("test_out.docx"))
    doc.save(input_path)
    
    detectors = [RegexDetector(), NERDetector(), ContextDetector()]
    resolver = EntityResolver()
    replacer = DeterministicReplacer()
    processor = DocumentProcessor(detectors, resolver, replacer)
    
    processor.process(input_path, output_path)
    
    out_doc = docx.Document(output_path)
    p_text = out_doc.paragraphs[0].text
    t_text = out_doc.tables[0].cell(0, 0).text
    
    assert "Alice Bobson" not in p_text
    assert "Alice Bobson" not in t_text
    assert p_text.count("are the promoters") == 1
    # Check deterministic replacement match
    # Since it's redacted in both places, the replacement should be identical
    repl_in_para = p_text.split(" and ")[0].strip()
    assert repl_in_para == t_text.strip()

def test_address_exclusion_from_global_memoization(tmpdir):
    doc = docx.Document()
    # Strong address context for 'Springfield'
    doc.add_paragraph("Our registered office is located at Springfield, IL 62704.")
    # Unrelated usage of Springfield as part of a facility name
    doc.add_paragraph("The Springfield Facility is operational.")
    
    input_path = str(tmpdir.join("test_in2.docx"))
    output_path = str(tmpdir.join("test_out2.docx"))
    doc.save(input_path)
    
    detectors = [RegexDetector(), NERDetector(), ContextDetector()]
    resolver = EntityResolver()
    replacer = DeterministicReplacer()
    processor = DocumentProcessor(detectors, resolver, replacer)
    
    processor.process(input_path, output_path)
    
    out_doc = docx.Document(output_path)
    p1 = out_doc.paragraphs[0].text
    p2 = out_doc.paragraphs[1].text
    
    # Address should be redacted in p1
    assert "Springfield, IL" not in p1
    # But Springfield should NOT be globally memoized and should survive in p2
    assert "Springfield Facility" in p2

def test_partial_substring_not_replaced(tmpdir):
    doc = docx.Document()
    doc.add_paragraph("John Smith is the CEO.")
    doc.add_paragraph("A blacksmith was working.")
    
    input_path = str(tmpdir.join("test_in3.docx"))
    output_path = str(tmpdir.join("test_out3.docx"))
    doc.save(input_path)
    
    detectors = [RegexDetector(), NERDetector(), ContextDetector()]
    resolver = EntityResolver()
    replacer = DeterministicReplacer()
    processor = DocumentProcessor(detectors, resolver, replacer)
    
    processor.process(input_path, output_path)
    
    out_doc = docx.Document(output_path)
    assert "John Smith" not in out_doc.paragraphs[0].text
    # "smith" inside "blacksmith" must NOT be redacted
    assert "blacksmith" in out_doc.paragraphs[1].text

def test_multi_run_preservation_with_memoization(tmpdir):
    doc = docx.Document()
    p = doc.add_paragraph("Jane Doe ")
    p.add_run("are our promoters.").bold = True
    
    p2 = doc.add_paragraph()
    p2.add_run("Jane ").bold = True
    p2.add_run("Doe").italic = True
    
    input_path = str(tmpdir.join("test_in4.docx"))
    output_path = str(tmpdir.join("test_out4.docx"))
    doc.save(input_path)
    
    detectors = [RegexDetector(), NERDetector(), ContextDetector()]
    resolver = EntityResolver()
    replacer = DeterministicReplacer()
    processor = DocumentProcessor(detectors, resolver, replacer)
    
    processor.process(input_path, output_path)
    
    out_doc = docx.Document(output_path)
    # The name should be fully redacted across runs
    assert "Jane" not in out_doc.paragraphs[1].text
    assert "Doe" not in out_doc.paragraphs[1].text
