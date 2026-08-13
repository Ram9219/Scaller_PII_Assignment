import os
import sys
import docx
import pytest

# Add parent to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from document.processor import DocumentProcessor
from core.resolver import EntityResolver
from core.replacement import DeterministicReplacer
from detectors import RegexDetector

def create_mock_docx(path):
    doc = docx.Document()
    doc.add_paragraph("This is a simple paragraph with an email: test@example.com")
    
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Header 1"
    table.rows[0].cells[1].text = "Header 2"
    table.rows[1].cells[0].text = "Data 1"
    table.rows[1].cells[1].text = "Data 2"
    
    # Add a section with header/footer
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "Header text"
    
    doc.save(path)

def test_docx_structure_preservation(tmpdir):
    input_path = str(tmpdir.join("input.docx"))
    output_path = str(tmpdir.join("output.docx"))
    
    create_mock_docx(input_path)
    
    # Process
    detectors = [RegexDetector()]
    processor = DocumentProcessor(detectors, EntityResolver(), DeterministicReplacer())
    processor.process(input_path, output_path)
    
    # Assertions
    doc_in = docx.Document(input_path)
    doc_out = docx.Document(output_path)
    
    # 1. Paragraph count
    assert len(doc_in.paragraphs) == len(doc_out.paragraphs), "Paragraph count mismatch"
    
    # 2. Table count
    assert len(doc_in.tables) == len(doc_out.tables), "Table count mismatch"
    
    # 3. Section count
    assert len(doc_in.sections) == len(doc_out.sections), "Section count mismatch"
    
    # 4. Table dimensions
    for i in range(len(doc_in.tables)):
        assert len(doc_in.tables[i].rows) == len(doc_out.tables[i].rows), f"Table {i} row count mismatch"
        assert len(doc_in.tables[i].columns) == len(doc_out.tables[i].columns), f"Table {i} column count mismatch"
        
    # 5. Header/Footer preservation check
    assert len(doc_in.sections[0].header.paragraphs) == len(doc_out.sections[0].header.paragraphs)
    assert doc_out.sections[0].header.paragraphs[0].text == "Header text"
    
    # 6. Formatting check on unchanged runs
    assert "This is a simple paragraph with an email:" in doc_out.paragraphs[0].text

def test_deterministic_replacement():
    """Verify that the same PII input yields the identical Faker replacement."""
    replacer1 = DeterministicReplacer()
    replacer2 = DeterministicReplacer()
    
    val1 = replacer1.get_replacement("rajesh.k@example.com", "EMAIL")
    val2 = replacer2.get_replacement("rajesh.k@example.com", "EMAIL")
    val3 = replacer1.get_replacement("rajesh.k@example.com", "EMAIL")
    
    # Must be deterministic across calls and across instances
    assert val1 == val2
    assert val1 == val3
    
    # A different email must yield a different value
    val4 = replacer1.get_replacement("different@example.com", "EMAIL")
    assert val1 != val4
