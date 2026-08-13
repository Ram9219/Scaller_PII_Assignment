import docx
from typing import Iterator, Tuple
from core.entities import EntityLocation

class DocumentTraverser:
    """
    Traverses a python-docx Document and yields (text, EntityLocation, paragraph_object).
    The paragraph_object is used later by the processor to apply replacements.
    """
    def __init__(self, doc: docx.Document):
        self.doc = doc
        self.seen_paragraphs = set()

    def _traverse_tables(self, tables, location_prefix) -> Iterator[Tuple[str, EntityLocation, object]]:
        for t_idx, table in enumerate(tables):
            headers = []
            if len(table.rows) > 0:
                for cell in table.rows[0].cells:
                    headers.append(cell.text.strip())

            for r_idx, row in enumerate(table.rows):
                is_header = (r_idx == 0)
                for c_idx, cell in enumerate(row.cells):
                    header_text = headers[c_idx] if c_idx < len(headers) else ""
                    
                    for p_idx, paragraph in enumerate(cell.paragraphs):
                        if paragraph._element in self.seen_paragraphs:
                            continue
                        self.seen_paragraphs.add(paragraph._element)
                        
                        text = paragraph.text
                        if text.strip():
                            loc_kwargs = dict(location_prefix)
                            loc_kwargs.update({
                                'table_index': t_idx, 
                                'row_index': r_idx, 
                                'cell_index': c_idx, 
                                'paragraph_index': p_idx,
                                'is_header': is_header
                            })
                            loc = EntityLocation(**loc_kwargs)
                            setattr(loc, 'header_text', header_text)
                            yield text, loc, paragraph

                    # Recursively traverse nested tables in the cell
                    yield from self._traverse_tables(cell.tables, location_prefix)

    def traverse(self) -> Iterator[Tuple[str, EntityLocation, object]]:
        self.seen_paragraphs = set() # Reset on new traversal
        
        # 1. Main Document Paragraphs
        for i, paragraph in enumerate(self.doc.paragraphs):
            if paragraph._element in self.seen_paragraphs:
                continue
            self.seen_paragraphs.add(paragraph._element)
            text = paragraph.text
            if text.strip():
                yield text, EntityLocation(paragraph_index=i), paragraph

        # 2. Main Tables
        yield from self._traverse_tables(self.doc.tables, {})

        # 3. Headers and Footers (from sections)
        for s_idx, section in enumerate(self.doc.sections):
            # Header paragraphs
            for p_idx, paragraph in enumerate(section.header.paragraphs):
                if paragraph._element in self.seen_paragraphs:
                    continue
                self.seen_paragraphs.add(paragraph._element)
                text = paragraph.text
                if text.strip():
                    yield text, EntityLocation(paragraph_index=p_idx, is_header=True), paragraph
                    
            # Header tables
            yield from self._traverse_tables(section.header.tables, {'is_header': True})
            
            # Footer paragraphs
            for p_idx, paragraph in enumerate(section.footer.paragraphs):
                if paragraph._element in self.seen_paragraphs:
                    continue
                self.seen_paragraphs.add(paragraph._element)
                text = paragraph.text
                if text.strip():
                    yield text, EntityLocation(paragraph_index=p_idx, is_footer=True), paragraph
                    
            # Footer tables
            yield from self._traverse_tables(section.footer.tables, {'is_footer': True})
