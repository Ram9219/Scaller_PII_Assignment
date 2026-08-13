import docx
import json

doc = docx.Document('../Red Herring Prospectus.docx')

samples = []

# Find some contact info (search for 'Email' or 'Contact' or 'Address')
for i, para in enumerate(doc.paragraphs):
    if len(samples) > 20: break
    text = para.text.strip()
    if ('email' in text.lower() or 'contact' in text.lower() or 'address' in text.lower() or 'director' in text.lower() or 'mr.' in text.lower() or 'date' in text.lower() or 'dob' in text.lower()):
        if 20 < len(text) < 500:
            samples.append({"text": text, "source": f"paragraph_{i}"})

# Find some interesting tables
for t_idx, table in enumerate(doc.tables):
    if len(samples) > 40: break
    if len(table.rows) > 0:
        header = [c.text.strip().replace('\n', ' ') for c in table.rows[0].cells]
        if any('name' in h.lower() or 'email' in h.lower() or 'address' in h.lower() or 'din' in h.lower() or 'dob' in h.lower() for h in header):
            for r_idx, row in enumerate(table.rows[1:3]): # just first couple rows
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip().replace('\n', ', ')
                    if text:
                        samples.append({
                            "text": text, 
                            "source": f"table_{t_idx}_row_{r_idx+1}_col_{c_idx}",
                            "header_context": header[c_idx] if c_idx < len(header) else ""
                        })

with open('raw_samples.json', 'w') as f:
    json.dump(samples, f, indent=4)
